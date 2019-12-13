import requests,os
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
from urllib.request import quote
from docx.shared import Inches
from PIL import Image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
def get_value():
    main_site = "http://www.chinaz.com/"
    text = requests.get(main_site).text
    soup = BeautifulSoup(text,'lxml')
    value = soup.find('input',{"name":"s"})['value']
    return value

def docuemt_wirte(arr,title):
    filename = "%s/%s.doc"%(str(datetime.now().date()),title)
    document = Document()
    for d in arr:
        print(d)
        if str(d).__contains__(".png") or str(d).__contains__(".jpg"):
            paragraph = document.add_paragraph()
            # 图片居中设置
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run("")
            im = Image.open(d)
            #基点数*914400
            #原图片尺寸 553x69 长宽比例为553/69=8.01
            #3*914400==word中7.62
            #2*914400==word中5.08
            px = 37.8006872852
            cm = 360000
            width = int(im.width / px * cm)
            height = int(im.height / px * cm)
            run.add_picture(d, width=width, height=height)
        else:
            document.add_paragraph(d)
    document.save(filename)
    print("file[%s] write done!!!"%filename)


def download_pic(url,filename):
    html = requests.get(url)
    with open(filename, "wb")as f:
        f.write(html.content)

def run():
    key="LOL"
    value = get_value()
    req_url="http://zhannei.baidu.com/cse/search?s=%s&q=%s&srt=lds&nsid=0&sti=1440"%(value,quote(key))
    print(req_url)
    text = requests.get(req_url).text
    soup = BeautifulSoup(text, 'lxml')
    arr = soup.find_all('div',{"class":"result f s0"})
    for a in arr:
        temp_soup = BeautifulSoup(str(a), 'lxml')
        article = temp_soup.find('a',{"cpos":"title"})
        article_url = str(article['href']).strip()
        print(article_url)
        res =requests.get(article_url).content
        text = str(res, 'utf-8')
        soup = BeautifulSoup(text, 'lxml')
        title = soup.find('h2',{"class":"title"}).text.replace("/","+")
        print(title)
        content = soup.find('div',{"id":"ctrlfscont"})
        temp_soup = BeautifulSoup(str(content),'lxml')
        p_arr = temp_soup.find_all('p')
        content_arr = []
        for p in p_arr:
            if str(p).__contains__("<img"):
                pic_url = str(p.img['src'])
                pic_name = pic_url.split("/")[-1].replace("/","+")
                local_file_path = "%s/picture/%s/"%(str(datetime.now().date()),title)
                if not os.path.exists(local_file_path):
                    os.makedirs(local_file_path)
                pic_path = os.path.join(local_file_path,pic_name)
                download_pic(pic_url,pic_path)
                basedir = os.path.abspath(os.path.dirname(__file__))
                ab_path  = os.path.join(basedir,pic_path).replace("\\","/")
                content_arr.append(ab_path)
            else:
                text_doc = str(p.text).replace("斗玩网(d.chinaz.com)原创：","")
                if not text_doc.__contains__(">>") and not text_doc.__contains__("点击下载") and not text_doc.__contains__("提取码") and not text_doc.__contains__("原贴地址") and not text_doc.__contains__("点击进入"):
                    content_arr.append(text_doc)
        docuemt_wirte(content_arr,title)




if __name__ == '__main__':
    run()