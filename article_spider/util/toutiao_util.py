from redis import StrictRedis
import os,requests,random
from datetime import datetime as dt
from docx import Document
from PIL import Image
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
#写入本地文件 工具类上级需要和引用文件上级同一目录
#请求工具
class RequestUtil():
    @staticmethod
    def get_header(cookie,referer=None):
        header = {
            'Connection': 'keep-alive',
            'Accept': 'application/json, text/plain, */*',
            'Accept-Encoding':'gzip, deflate, br',
            'Accept-Language': 'zh-CN,zh;q=0.9',
            'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/72.0.3626.64 Safari/537.36',
            'Host':'mp.toutiao.com',
            'Cookie': cookie
        }
        if referer is not None:
            header["Referer"] = referer
        return header


def write_log(line):
    file_name = "G:/code/Article/logs/%s.log"%str(dt.now().date())
    with open(file_name,'a+',encoding='utf-8') as f:
        f.write("[%s]%s\n"%(str(dt.now().strftime('%Y-%m-%d %H:%M:%S')),line))
        print("[%s]%s"%(str(dt.now().strftime('%Y-%m-%d %H:%M:%S')),line))

if __name__ == '__main__':
    write_log('test')

def tag_filter(content):
    tag_list = ['strong']
    for tag in tag_list:
        prefix = "<%s>"%tag
        suffix = "<%s/>"%tag
        if content.__contains__(prefix) or content.__contains__(suffix):
            content = content.replace(prefix,"").replace(suffix,"")
    return content

#可能包含原作者信息的 需要从文章内容中过滤
def content_filter(content):
    words=['我是']
    for w in words:
        if content.__contains__(w):
            if content.__contains__(','):
                content_arr = content.split(',')
                arr = [x for x in content_arr if not str(x).__contains__(w)]
                content=",".join(arr)
            elif content.__contains__('，'):
                content_arr = content.split('，')
                arr = [x for x in content_arr if not str(x).__contains__(w)]
                content = "，".join(arr)
            else:
                content = ""
    return content



def download_pic(url,filename):
    html = requests.get(url)
    with open(filename, "wb")as f:
        f.write(html.content)


def convert_doc(content_list,title):
    # 上级目录-source
    upper = os.path.abspath(os.path.join(os.getcwd(), ".."))
    source = os.path.join(upper, 'source').replace("\\", "/")
    if not os.path.exists(source):
        os.mkdir(source)

    filename = "%s/%s/%s.doc" % (source,str(dt.now().date()), title)
    document = Document()
    for d in content_list:
        print(d)
        if str(d).__contains__(".png") or str(d).__contains__(".jpg"):
            paragraph = document.add_paragraph()
            # 图片居中设置
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run("")
            im = Image.open(d)
            # 基点数*914400
            # 原图片尺寸 553x69 长宽比例为553/69=8.01
            # 3*914400==word中7.62
            # 2*914400==word中5.08
            px = 37.8006872852
            cm = 360000
            width = int(im.width / px * cm)
            height = int(im.height / px * cm)
            run.add_picture(d, width=width, height=height)
        else:
            document.add_paragraph(d)
    document.save(filename)
    write_log("file[%s] write done!!!" % filename)

def format_p(p_arr,title):
    content_list = []
    for p in p_arr:
        if str(p).__contains__("img"):
            pic_url = str(p.find('img')['src'])
            if not pic_url.startswith("http"):
               pic_url = "http:"+pic_url
            pic_name = "".join(random.sample('zyxwvutsrqponmlkjihgfedcba',20))+ ".png"
            # 上级目录
            upper = os.path.abspath(os.path.join(os.getcwd(), ".."))
            source = os.path.join(upper, 'source').replace("\\", "/")
            if not os.path.exists(source):
                os.mkdir(source)

            local_file_path = "%s/%s/picture/%s/" % (source, str(dt.now().date()), title)
            if not os.path.exists(local_file_path):
                os.makedirs(local_file_path)
            pic_path = os.path.join(local_file_path, pic_name)
            download_pic(pic_url, pic_path)
            content = pic_path
        else:
            content = str(p.text).strip()
        content = tag_filter(str(content))
        content = content_filter(content)
        write_log(content)
        if content!='':
            content_list.append(content)
    return content_list